from pypdf import PdfReader
from docx import Document
from pathlib import Path


BASE_DIR = Path(__file__).resolve().parent.parent.parent

resume_path = BASE_DIR / "resume" / "kumargaurav.pdf"


if resume_path.suffix.lower() not in [".pdf",".docx"]:
    raise ValueError ("only pdf or docx files are allowed")

def read_resume(resume_path):
    if resume_path.suffix.lower() == ".pdf":
        return read_pdf(resume_path)
    elif resume_path.suffix.lower() == ".docx":
        return read_docx(resume_path)
    else:
        return None

def read_pdf(resume_path):
    reader = PdfReader(resume_path)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text

def read_docx(resume_path):
    document = Document(resume_path)
    text = ""
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + "\n"
    
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + "\n"
    return text


